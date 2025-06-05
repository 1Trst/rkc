#!/usr/bin/env python3
"""
Streamlit Article Cleaner and Newsletter Generator: upload, clean, generate, and download Word documents
Enhanced version with Newsdesk support and image handling
"""
import streamlit as st
import zipfile
import io
import tempfile
import os
import re
import base64
import unicodedata
from datetime import datetime
from docx import Document as DocxDocument
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openai import AzureOpenAI
import asyncio
import httpx
import json
from concurrent.futures import ThreadPoolExecutor
from PIL import Image
import fitz  # PyMuPDF for PDF processing with images

# Azure settings - Get from environment variables
key = os.getenv("AZURE_OPENAI_API_KEY")
AZURE_OPENAI_ENDPOINT = os.getenv("AZURE_OPENAI_ENDPOINT", "https://rkcazureai.cognitiveservices.azure.com/")

# Language selection
LANGUAGES = {
    "Français": "French",
    "English": "English", 
    "Deutsch": "German"
}

def check_password():
    """Returns True if the user has entered the correct password."""
    
    def password_entered():
        """Checks whether a password entered by the user is correct."""
        if st.session_state["password"] == "rkcnewsletterWS2025":
            st.session_state["password_correct"] = True
            del st.session_state["password"]  # Don't store the password
        else:
            st.session_state["password_correct"] = False

    # Return True if password is validated
    if st.session_state.get("password_correct", False):
        return True

    # Show input for password
    st.markdown("### 🔐 Authentication Required")
    st.text_input(
        "Please enter the password to access the application:",
        type="password",
        on_change=password_entered,
        key="password",
        placeholder="Enter password..."
    )
    
    if "password_correct" in st.session_state:
        st.error("😞 Password incorrect. Please try again.")
    
    st.markdown("---")
    st.info("💡 Contact your administrator if you need access.")
    return False

# Check password before showing the main app
if not check_password():
    st.stop()

# UI
st.title("📝 Article Cleaner & Newsletter Generator")

# Show API key status in sidebar
with st.sidebar:
    if key:
        st.success("✅ Azure OpenAI API Key loaded")
        st.info(f"🔗 Endpoint: {AZURE_OPENAI_ENDPOINT}")
    else:
        st.error("❌ Azure OpenAI API Key not found")
        st.warning("Please set AZURE_OPENAI_API_KEY environment variable")

# Stop execution if no API key
if not key:
    st.error("🔑 **Configuration Error**: Azure OpenAI API Key not found in environment variables.")
    st.markdown("""
    **For local development:**
    1. Create a `.env` file in your project directory
    2. Add: `AZURE_OPENAI_API_KEY=your_api_key_here`
    
    **For deployment:**
    - Set the `AZURE_OPENAI_API_KEY` environment variable in your deployment platform
    """)
    st.stop()

# Main tabs
tab1, tab2, tab3 = st.tabs(["🧹 Article Cleaning", "📰 Newsletter Generator", "⚙️ Settings"])

# Initialize Azure OpenAI client
def get_client(key):
    """Get Azure OpenAI client properly"""
    if not key:
        return None
    try:
        from openai import AzureOpenAI
        
        client = AzureOpenAI(
            api_key=key,
            api_version="2024-02-15-preview",
            azure_endpoint=AZURE_OPENAI_ENDPOINT,
        )
        return client
    except Exception as e:
        st.error(f"Error initializing Azure OpenAI client: {str(e)}")
        return None

# Ensure Templates directory exists
def ensure_templates_dir():
    """Ensure Templates directory exists with at least a Generic template"""
    os.makedirs("Templates", exist_ok=True)
    # Create a generic template if none exists
    generic_path = os.path.join("Templates", "2025_Generic_Template.docx")
    if not os.path.exists(generic_path):
        doc = DocxDocument()
        doc.save(generic_path)
    return os.path.exists("Templates")

# Enhanced file handling functions with image extraction
def extract_images_from_pdf(file_path, workdir):
    """Extract images from PDF and return image info - improved for better filtering"""
    images_info = []
    try:
        doc = fitz.open(file_path)
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            image_list = page.get_images(full=True)
            
            for img_index, img in enumerate(image_list):
                xref = img[0]
                pix = fitz.Pixmap(doc, xref)
                
                # Filter out very small images (likely logos/icons) and ensure proper format
                if pix.n - pix.alpha < 4 and pix.width > 100 and pix.height > 100:  # Minimum size filter
                    img_filename = f"page_{page_num+1}_img_{img_index+1}.png"
                    img_path = os.path.join(workdir, img_filename)
                    
                    # Convert to RGB if needed and save
                    if pix.n == 4:  # CMYK
                        pix = fitz.Pixmap(fitz.csRGB, pix)
                    
                    pix.save(img_path)
                    
                    images_info.append({
                        'page': page_num + 1,
                        'filename': img_filename,
                        'path': img_path,
                        'description': f"Content image from page {page_num + 1}",
                        'width': pix.width,
                        'height': pix.height
                    })
                
                pix = None
        doc.close()
    except Exception as e:
        st.warning(f"Could not extract images from PDF: {str(e)}")
    
    return images_info

def clean_source_name(source):
    """Clean source name by removing 'Online' except for specific cases"""
    if not source:
        return "Unknown"
    
    # Exceptions: keep "Online" for these sources
    exceptions = ["Wallstreet Online", "KMA Online"]
    
    # Check if the source is in exceptions (case-insensitive)
    for exception in exceptions:
        if source.lower() == exception.lower():
            return source
    
    # Remove "Online" from other sources
    if source.endswith(" Online"):
        return source[:-7]  # Remove " Online"
    
    return source

def clean_title_for_filename(title):
    """Clean title for use in filename by removing only problematic characters"""
    if not title:
        return "Untitled"
    
    # Only replace characters that are actually problematic for Windows/filesystem
    # These are the forbidden characters: < > : " / \ | ? *
    invalid_chars = {
        '<': '',   # Remove
        '>': '',   # Remove  
        ':': ' ',  # Replace with space
        '"': "'",  # Replace with single quote
        '/': ' ',  # Replace with space
        '\\': ' ', # Replace with space
        '|': ' ',  # Replace with space
        '?': '',   # Remove
        '*': '',   # Remove
    }
    
    # Apply replacements
    for old, new in invalid_chars.items():
        title = title.replace(old, new)
    
    # Clean up multiple spaces
    title = re.sub(r'\s+', ' ', title)
    title = title.strip()
    
    # Limit length
    if len(title) > 100:
        title = title[:100].rstrip()
    
    return title or "Untitled"

def format_date_french(date_str):
    """Convert date to ddmmmyyyy format with French months"""
    if not date_str or date_str == "No date":
        now = datetime.now()
        return format_date_french_from_datetime(now)
    
    french_months = {
        1: 'jan', 2: 'fev', 3: 'mar', 4: 'avr',
        5: 'mai', 6: 'jui', 7: 'jul', 8: 'aou',
        9: 'sep', 10: 'oct', 11: 'nov', 12: 'dec'
    }
    
    # Try to parse various date formats
    date_formats = [
        "%d %B %Y", "%d %b %Y", "%Y-%m-%d", "%d/%m/%Y",
        "%d-%m-%Y", "%d.%m.%Y", "%d %m %Y", "%B %d, %Y", "%b %d, %Y"
    ]
    
    parsed_date = None
    for fmt in date_formats:
        try:
            parsed_date = datetime.strptime(date_str.strip(), fmt)
            break
        except ValueError:
            continue
    
    if not parsed_date:
        # Extract numbers and try to parse
        numbers = re.findall(r'\d+', date_str)
        if len(numbers) >= 3:
            try:
                if int(numbers[0]) > 31:  # Year first
                    year, month, day = int(numbers[0]), int(numbers[1]), int(numbers[2])
                else:  # Day first
                    day, month, year = int(numbers[0]), int(numbers[1]), int(numbers[2])
                
                if year < 100:
                    year += 2000 if year < 50 else 1900
                
                parsed_date = datetime(year, month, day)
            except (ValueError, IndexError):
                pass
    
    if parsed_date:
        return format_date_french_from_datetime(parsed_date)
    else:
        return format_date_french_from_datetime(datetime.now())

def format_date_french_from_datetime(dt):
    """Format datetime object to ddmmmyyyy with French months"""
    french_months = {
        1: 'jan', 2: 'fev', 3: 'mar', 4: 'avr',
        5: 'mai', 6: 'jui', 7: 'jul', 8: 'aou',
        9: 'sep', 10: 'oct', 11: 'nov', 12: 'dec'
    }
    
    day = f"{dt.day:02d}"
    month = french_months[dt.month]
    year = str(dt.year)
    
    return f"{day}{month}{year}"

def generate_docx_filename(source, title, date):
    """Generate DOCX filename in format: Source_Title_date.docx"""
    clean_source = clean_source_name(source)
    clean_title = clean_title_for_filename(title)
    formatted_date = format_date_french(date)
    
    # Construct filename - only use underscores to separate main components
    filename = f"{clean_source}_{clean_title}_{formatted_date}.docx"
    
    return filename

def read_file_content_with_images(file_path, workdir):
    """Read file content and extract images if available - improved text extraction"""
    ext = file_path.split('.')[-1].lower()
    content = ""
    images = []
    
    try:
        if ext == 'pdf':
            # Extract images first
            images = extract_images_from_pdf(file_path, workdir)
            
            # Enhanced text extraction using PyMuPDF for better formatting
            try:
                import fitz
                doc = fitz.open(file_path)
                text_blocks = []
                
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    # Get text with better formatting preservation
                    page_text = page.get_text("text")
                    if page_text.strip():
                        text_blocks.append(page_text)
                
                content = "\n\n".join(text_blocks)
                doc.close()
            except:
                # Fallback to PyPDF2 if fitz fails
                import PyPDF2
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    content = "\n".join([page.extract_text() for page in reader.pages])
                
        elif ext == 'txt':
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
        elif ext == 'docx':
            doc = DocxDocument(file_path)
            content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        elif ext == 'md':
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
        else:
            content = ""
            
    except Exception as e:
        st.error(f"Error reading {file_path}: {str(e)}")
        
    return content, images

# Helper function for reading file content (backward compatibility)
def read_file_content(file_path):
    """Read file content based on extension (without images)"""
    ext = file_path.split('.')[-1].lower()
    
    try:
        if ext == 'pdf':
            import PyPDF2
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                return "\n".join([page.extract_text() for page in reader.pages])
        elif ext == 'txt':
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        elif ext == 'docx':
            doc = DocxDocument(file_path)
            return "\n".join([paragraph.text for paragraph in doc.paragraphs])
        elif ext == 'md':
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        else:
            return ""
    except Exception as e:
        st.error(f"Error reading {file_path}: {str(e)}")
        return ""

# Get language-specific prompts
def get_language_instructions(language, article_summary_sentences=3, exec_summary_sentences=7):
    """Get language-specific instructions for prompts"""
    return {
        "summary_lang": language.upper(),
        "summary_instruction": f"Generate around {article_summary_sentences} sentences summary in {language.upper()}",
        "exec_summary_intro": f"You are a consultant in a French consulting firm. Your task is to write an Executive summary of around {exec_summary_sentences} sentences based on the following articles:",
        "exec_summary_requirements": f"""This summary should:
1. Synthesize the key ideas and points from the articles
2. Identify trends and points of divergence
3. Highlight relevant implications for our consulting firm (no vague or general implications)
4. If relevant, formulate concrete recommendations

Respond only with the executive summary in {language.upper()} (around {exec_summary_sentences} sentences), without additional introduction or conclusion."""
    }

# Enhanced prompts for different article types
def get_newsdesk_prompt(text, language="French", article_summary_sentences=3):
    """Get specialized prompt for Nexis Newsdesk articles"""
    lang_instructions = get_language_instructions(language, article_summary_sentences)
    
    if language == "French":
        translation_instruction = "Keep the title and article text in their original language if they are already in French, otherwise translate them to French."
        output_language = "French"
    elif language == "English":
        translation_instruction = "Translate the title and article text to English."
        output_language = "English"
    else:  # German
        translation_instruction = "Translate the title and article text to German."
        output_language = "German"

    return f"""You are processing a Nexis Newsdesk document that may contain one or more articles from various sources.

CRITICAL OUTPUT FORMAT: Return ONLY a JSON array starting with [ and ending with ]. Do NOT wrap in any other object.

ARTICLE IDENTIFICATION PATTERN:
Nexis Newsdesk documents follow this structure:
1. Header information to IGNORE (Exclusion de responsabilité, Flux, Plage de dates, Téléchargé)
2. Article titles in blue/bold formatting
3. Source line format: "Source Name | Author | Date Time"
4. Full article content
5. Sometimes footer information to IGNORE

EXTRACTION PROCESS:
1. IDENTIFY each individual article by looking for:
   - Clear article titles (often in blue/larger text)
   - Source attribution lines (Source | Author | Date pattern)
   - Article content that follows

2. For EACH article found:
   - Title: Extract the exact headline
   - Source: Extract from source line - common sources include: "Le Figaro Online", "Les Echos", "Les Echos Investir", "Le Monde", "Financial Times", "Handelsblatt", "La Lettre", "Consultor", "La Correspondance économique"
   - Date: Extract from source line (formats like "DD MMM YYYY" or "DD avr. YYYY")
   - Content: Full article text, removing only Nexis headers/footers BUT PRESERVING ORIGINAL PARAGRAPH STRUCTURE
   - Summary: Generate {article_summary_sentences} sentences in {output_language}

3. REMOVE these Nexis elements:
   - "Exclusion de responsabilité : Ce document Nexis Newsdesk®..."
   - "Flux: [anything]"
   - "Plage de dates: [anything]"
   - "Téléchargé(e): [date] par [email]"
   - Footer copyright notices
   - Navigation elements

CRITICAL - PARAGRAPH PRESERVATION:
- PRESERVE the original paragraph structure and line breaks from the source articles
- Do NOT merge multiple paragraphs into one large block of text
- Maintain natural paragraph separations as they appear in the original
- Keep the logical flow and structure of the original article
- Each paragraph should remain as a separate paragraph in the cleaned_text
- Use proper line breaks (\\n\\n) to separate paragraphs in the JSON string

REQUIRED OUTPUT FORMAT (must be a direct array):
[
    {{
        "title": "exact article title without special characters",
        "source": "source name from source line",
        "date": "DD MMM YYYY",
        "year": "YYYY",
        "cleaned_text": "full article content with PRESERVED paragraph structure and line breaks",
        "summary": "{article_summary_sentences} sentence summary in {output_language}"
    }},
    {{
        "title": "second article title if exists",
        "source": "second source",
        "date": "DD MMM YYYY",
        "year": "YYYY", 
        "cleaned_text": "second article content with PRESERVED paragraph structure",
        "summary": "summary of second article"
    }}
]

IMPORTANT RULES:
- Start response with [
- End response with ]
- Do NOT use {{"articles": [...] }} wrapper format
- Extract ALL articles found in the document
- If only 1 article found, return array with 1 object
- If multiple articles found, return array with multiple objects
- Each article object must have all 6 fields: title, source, date, year, cleaned_text, summary
- PRESERVE paragraph structure in cleaned_text - do NOT create one big text block
- Use \\n\\n to separate paragraphs in the cleaned_text field
- {translation_instruction}

Document to process:
{text}"""

# Process articles with Azure OpenAI with retry logic
async def process_article_async(client, text, language="French", article_summary_sentences=3, is_newsdesk=False):
    """Process a single article with Azure OpenAI with retry logic"""
    import time
    
    if is_newsdesk:
        prompt = get_newsdesk_prompt(text, language, article_summary_sentences)
    else:
        lang_instructions = get_language_instructions(language, article_summary_sentences)
        
        # Get language-specific instructions for translation
        if language == "French":
            translation_instruction = "Keep the title and article text in their original language if they are already in French, otherwise translate them to French."
            output_language = "French"
        elif language == "English":
            translation_instruction = "Translate the title and article text to English."
            output_language = "English"
        else:  # German
            translation_instruction = "Translate the title and article text to German."
            output_language = "German"
        
        # Standard prompt for regular articles
        prompt = f"""You will be provided a poorly copy/pasted article from a single journal/website. Please extract the following information and clean this article in JSON format:

1. Extract and clean the title (remove problematic characters like &,/,<,>,#,»,«), then translate it to {output_language}.
2. Extract the source of the article (journal or website name) it has to be in the following list : ["Consultor","Financial Times","Handelsblatt","La Lettre_du_Conseil","La Lettre","Les Echos Investir","Les Echos","Le Monde","Le Figaro Online","La Correspondance économique"]. Pay close attention to this part, and detect the difference between "Les Echos" et "Les Echos Investir"
3. Extract the date (format: d MMMM yyyy)
4. Clean the article by removing any website boilerplate, ads, or irrelevant content. IMPORTANT: Please try to respect and understand the different paragraphs of the original article and reproduce those in your cleaned text - PRESERVE the original paragraph structure and line breaks. Do NOT merge multiple paragraphs into one large block of text. Then translate the entire cleaned article text to {output_language}.
5. {lang_instructions["summary_instruction"]}
6. Double check the source : is what you found correct ?

IMPORTANT: {translation_instruction}

CRITICAL - PARAGRAPH PRESERVATION:
- PRESERVE the original paragraph structure and line breaks from the source article
- Do NOT merge multiple paragraphs into one large block of text
- Maintain natural paragraph separations as they appear in the original
- Keep the logical flow and structure of the original article
- Each paragraph should remain as a separate paragraph in the cleaned_text
- Use proper line breaks (\\n\\n) to separate paragraphs in the JSON string

Return ONLY a JSON object with these fields:
{{
    "title": "cleaned title translated to {output_language}",
    "source": "source (keep original name)", 
    "date": "date",
    "year": "yyyy",
    "cleaned_text": "full cleaned article text translated to {output_language} with PRESERVED paragraph structure and proper \\n\\n separators",
    "summary": "around {article_summary_sentences} sentences summary in {lang_instructions['summary_lang']}"
}}

Here's the text to process:

{text}"""

    max_retries = 3
    for attempt in range(max_retries):
        try:
            response = await asyncio.to_thread(
                client.chat.completions.create,
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                response_format={"type": "json_object"},
                temperature=0
            )
            
            return response.choices[0].message.content
            
        except Exception as e:
            error_str = str(e)
            if "429" in error_str or "rate limit" in error_str.lower():
                if attempt < max_retries - 1:
                    wait_time = 60 * (attempt + 1)  # Wait 60, 120, 180 seconds
                    st.warning(f"⏳ Rate limit hit. Waiting {wait_time} seconds before retry {attempt + 2}/{max_retries}...")
                    time.sleep(wait_time)
                    continue
                else:
                    st.error(f"❌ Rate limit exceeded after {max_retries} attempts. Please upgrade your Azure OpenAI tier or try again later.")
                    return "{}" if not is_newsdesk else "[]"
            else:
                st.error(f"Error processing article: {str(e)}")
                return "{}" if not is_newsdesk else "[]"

# Generate executive summary
async def generate_executive_summary(client, articles, language="French", exec_summary_sentences=7):
    """Generate an executive summary from article summaries"""
    lang_instructions = get_language_instructions(language, 3, exec_summary_sentences)
    
    # Create a prompt with all article summaries
    article_summaries = []
    for article in articles:
        article_summaries.append(f"**{article['title']}**: {article['summary']}")
    
    prompt = f"""{lang_instructions["exec_summary_intro"]}

{chr(10).join(article_summaries)}

{lang_instructions["exec_summary_requirements"]}"""

    try:
        response = await asyncio.to_thread(
            client.chat.completions.create,
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0
        )
        
        return response.choices[0].message.content
    except Exception as e:
        st.error(f"Error generating executive summary: {str(e)}")
        return "Error generating executive summary."

# Process multiple articles in parallel
async def process_articles(client, articles, language="French", article_summary_sentences=3, is_newsdesk=False):
    """Process multiple articles concurrently"""
    tasks = [process_article_async(client, text, language, article_summary_sentences, is_newsdesk) for text in articles]
    return await asyncio.gather(*tasks)

# Enhanced template handling with new naming convention
def get_template_path(source, is_newsdesk=False):
    """Get template path with new 2025_source_Template.docx naming convention (keeping spaces in source name)"""
    # For Newsdesk mode, normalize "Source Online" to "Source"
    if is_newsdesk and source.endswith(" Online"):
        source = source.replace(" Online", "")
    
    # Keep spaces in source name for template filename
    template_name = f"2025_{source}_Template.docx"
    template_path = os.path.join('Templates', template_name)
    
    # Fall back to generic template if source-specific doesn't exist
    if not os.path.exists(template_path):
        generic_path = os.path.join('Templates', '2025_Generic_Template.docx')
        if os.path.exists(generic_path):
            return generic_path
        else:
            # Create generic template if it doesn't exist
            doc = DocxDocument()
            doc.save(generic_path)
            return generic_path
    
    return template_path

# Enhanced Word document creation with image support and new template system
def create_word_doc_with_images(article_data, output_path, images_info=None, is_newsdesk=False):
    """Create a Word document from cleaned article data with images using new template system"""
    source = article_data.get("source", "Generic")
    
    # Use new template system with newsdesk awareness
    template_path = get_template_path(source, is_newsdesk)
    doc = DocxDocument(template_path)
    
    # CLEAR ALL EXISTING CONTENT FROM TEMPLATE
    # Remove all paragraphs from the document
    for paragraph in doc.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None
    
    # Generate proper filename
    source_name = article_data.get("source", "Unknown")
    title = article_data.get("title", "Untitled")
    date = article_data.get("date", "No date")
    
    proper_filename = generate_docx_filename(source_name, title, date)
    
    # Update the output path to use the proper filename
    output_dir = os.path.dirname(output_path)
    new_output_path = os.path.join(output_dir, proper_filename)
    
    # Set document properties using the cleaned filename (without .docx)
    core_props = doc.core_properties
    core_props.title = proper_filename[:-5]  # Remove .docx extension
    core_props.author = source_name
    
    # Add title
    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(f"{title} — {source_name}, {date}")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.name = 'Aptos'
    
    # Add images if available
    if images_info:
        for img_info in images_info:
            try:
                img_paragraph = doc.add_paragraph()
                img_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                run = img_paragraph.add_run()
                run.add_picture(img_info['path'], width=Inches(5))
                
                # Add image caption
                caption = doc.add_paragraph(f"Figure: {img_info['description']}")
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_run = caption.runs[0]
                caption_run.italic = True
                caption_run.font.size = Pt(10)
                
            except Exception as e:
                st.warning(f"Could not add image {img_info['filename']}: {str(e)}")
    
    # Add article body - SPLIT INTO SEPARATE PARAGRAPHS TO FIX JUSTIFICATION
    article_text = article_data.get("cleaned_text", "No content available")
    
    # Split the text by double line breaks (paragraph separators)
    paragraphs = article_text.split('\n\n')
    
    # Add each paragraph separately
    for paragraph_text in paragraphs:
        if paragraph_text.strip():  # Only add non-empty paragraphs
            # Clean up single line breaks within the paragraph (replace with spaces)
            clean_paragraph = paragraph_text.replace('\n', ' ').strip()
            
            # Add the paragraph
            body_paragraph = doc.add_paragraph(clean_paragraph)
            body_paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Optional: Add some spacing between paragraphs
            body_paragraph.paragraph_format.space_after = Pt(6)
    
    # Save document with the proper filename
    doc.save(new_output_path)
    return new_output_path


# Create newsletter document with images
def create_newsletter_doc_with_images(exec_summary, articles, output_path, all_images=None):
    """Create a Word document with newsletter content and images"""    
    doc = DocxDocument()
    
    # Set document properties
    core_props = doc.core_properties
    core_props.title = "AI Generated Newsletter"
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title.add_run("AI Generated Newsletter")
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.name = 'Aptos'
    
    # Date
    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.LEFT
    from datetime import datetime
    date_run = date.add_run(datetime.now().strftime("%d %B %Y"))
    date_run.italic = True
    date_run.font.size = Pt(12)
    
    # Executive summary 
    doc.add_paragraph()
    exec_heading = doc.add_paragraph()
    exec_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    exec_heading_run = exec_heading.add_run("Executive Summary")
    exec_heading_run.bold = True
    exec_heading_run.font.size = Pt(16)    
    
    # Executive summary content - SPLIT INTO PARAGRAPHS
    exec_paragraphs = exec_summary.split('\n\n')
    for paragraph_text in exec_paragraphs:
        if paragraph_text.strip():
            clean_paragraph = paragraph_text.replace('\n', ' ').strip()
            exec_content = doc.add_paragraph(clean_paragraph)
            exec_content.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            exec_content.paragraph_format.space_after = Pt(6)
    
    # Add sample images if available
    if all_images:
        doc.add_paragraph()
        img_heading = doc.add_paragraph()
        img_heading_run = img_heading.add_run("Visual Content")
        img_heading_run.bold = True
        img_heading_run.font.size = Pt(14)
        
        for i, img_info in enumerate(all_images[:3]):  # Limit to first 3 images
            try:
                img_paragraph = doc.add_paragraph()
                img_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                run = img_paragraph.add_run()
                run.add_picture(img_info['path'], width=Inches(4))
                
                caption = doc.add_paragraph(f"Figure {i+1}: {img_info['description']}")
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_run = caption.runs[0]
                caption_run.italic = True
                caption_run.font.size = Pt(10)
                
            except Exception as e:
                st.warning(f"Could not add image to newsletter: {str(e)}")
    
    # Separator
    doc.add_paragraph()
    separator = doc.add_paragraph("* * *")
    separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    # Add articles
    for article in articles:
        # Article title
        article_title = doc.add_paragraph()
        article_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title_text = f"{article['title']} — {article['source']}, {article['date']}"
        title_run = article_title.add_run(title_text)
        title_run.bold = True
        title_run.font.size = Pt(14)
        
        # Article summary - SPLIT INTO SEPARATE PARAGRAPHS TO FIX FORMATTING
        summary_text = article['summary']
        summary_paragraphs = summary_text.split('\n\n')
        
        for paragraph_text in summary_paragraphs:
            if paragraph_text.strip():
                # Clean up single line breaks within the paragraph
                clean_paragraph = paragraph_text.replace('\n', ' ').strip()
                
                # Add the paragraph
                summary_paragraph = doc.add_paragraph(clean_paragraph)
                summary_paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                summary_paragraph.paragraph_format.space_after = Pt(6)
        
        # Separator between articles
        doc.add_paragraph()
        separator = doc.add_paragraph("* * *")
        separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
    
    # Save document
    doc.save(output_path)
    return output_path

# Create download link for file
def get_download_link(file_path, link_text):
    """Generate a download link for a file"""
    with open(file_path, "rb") as f:
        data = f.read()
    b64 = base64.b64encode(data).decode()
    file_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    href = f'<a href="data:{file_type};base64,{b64}" download="{os.path.basename(file_path)}">{link_text}</a>'
    return href

# Create zip file from multiple files
def create_zip_from_files(file_paths):
    """Create a zip file from multiple files"""
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w") as zf:
        for file_path in file_paths:
            zf.write(file_path, os.path.basename(file_path))
    return zip_buffer.getvalue()

# Extract file paths with temporary directory for images
def get_paths_with_workdir(files):
    """Get file paths from uploaded files with a shared working directory"""
    workdir = tempfile.mkdtemp()
    paths = []
    for f in files or []:
        dest = os.path.join(workdir, f.name)
        with open(dest, 'wb') as out:
            out.write(f.read())
        paths.append(dest)
    return paths, workdir

# Settings tab
with tab3:
    st.header("⚙️ Settings")
    
    # Language selection
    selected_language = st.selectbox(
        "Choose output language / Choisir la langue de sortie:",
        options=list(LANGUAGES.keys()),
        index=0,  # Default to French
        help="This will affect the language of article summaries and newsletter content",
        key="language_selector"
    )
    
    # Store in session state
    st.session_state.output_language = LANGUAGES[selected_language]
    
    st.info(f"Selected language: {selected_language} / Langue sélectionnée: {selected_language}")
    
    # Summary length settings
    st.subheader("📏 Summary Length Settings")
    
    col1, col2 = st.columns(2)
    
    with col1:
        article_summary_sentences = st.number_input(
            "Article summary length (sentences):",
            min_value=1,
            max_value=10,
            value=3,
            step=1,
            help="Number of sentences for each article summary in the newsletter",
            key="article_summary_length"
        )
        st.session_state.article_summary_sentences = article_summary_sentences
    
    with col2:
        exec_summary_sentences = st.number_input(
            "Executive summary length (sentences):",
            min_value=3,
            max_value=15,
            value=7,
            step=1,
            help="Number of sentences for the executive summary",
            key="exec_summary_length"
        )
        st.session_state.exec_summary_sentences = exec_summary_sentences
    
    # Preview of settings
    st.subheader("📋 Current Settings Preview")
    st.info(f"""
    **Language:** {selected_language}
    **Article summaries:** Around {article_summary_sentences} sentences each
    **Executive summary:** Around {exec_summary_sentences} sentences
    """)
    
    st.markdown("---")
    st.markdown("💡 **Tip:** These settings will be applied to all new articles and newsletters you generate.")

# Article Cleaning Tab
with tab1:
    st.header("🧹 Article Cleaning")
    
    # Check if Templates directory exists
    ensure_templates_dir()
    
    # Get language setting
    output_language = st.session_state.get('output_language', 'French')
    article_summary_sentences = st.session_state.get('article_summary_sentences', 3)
    
    st.subheader("Processing Options")
    
    # Add Newsdesk toggle
    col1, col2 = st.columns([1, 1])
    with col1:
        is_newsdesk = st.toggle(
            "📰 Newsdesk Mode",
            value=False,
            help="Enable this for Nexis Newsdesk documents that contain multiple articles"
        )
    
    with col2:
        include_images = st.toggle(
            "🖼️ Include Images",
            value=True,
            help="Extract and include images in the final documents (works with PDFs)"
        )
    
    if is_newsdesk:
        st.info("🔧 **Newsdesk Mode**: Optimized for processing Nexis Newsdesk documents with multiple articles")
    
    st.subheader("Input Methods")
    
    # Input method selection
    input_method = st.radio(
        "Choose input method:",
        ["Upload Files", "Copy-Paste Text"],
        horizontal=True
    )
    
    articles_to_process = []
    article_names = []
    all_images_info = []
    workdir = None
    
    if input_method == "Upload Files":
        # File upload
        uploaded_files = st.file_uploader(
            "Upload one or multiple files", 
            type=["pdf", "txt", "docx", "md"], 
            accept_multiple_files=True
        )
        
        if uploaded_files:
            paths, workdir = get_paths_with_workdir(uploaded_files)
            for path in paths:
                if include_images:
                    content, images = read_file_content_with_images(path, workdir)
                    all_images_info.extend(images)
                else:
                    content = read_file_content(path)
                    
                if content:
                    articles_to_process.append(content)
                    article_names.append(os.path.basename(path))
    
    else:  # Copy-Paste Text
        # Text input
        pasted_text = st.text_area(
            "Paste your article text here:",
            height=300,
            placeholder="Copy and paste your article content here..."
        )
        
        if pasted_text.strip():
            articles_to_process.append(pasted_text.strip())
            article_names.append("Pasted_Article")
    
    # Show current settings and image info
    if articles_to_process:
        col1, col2 = st.columns([2, 1])
        with col1:
            st.info(f"Ready to process {len(articles_to_process)} articles with {article_summary_sentences} sentences summaries in {output_language}")
        with col2:
            if all_images_info:
                st.success(f"🖼️ Found {len(all_images_info)} images")
    
    # Clean articles button
    if st.button("🧹 Clean Articles", use_container_width=True):
        client = get_client(key)
        
        if not client:
            st.error("Failed to initialize Azure OpenAI client")
        elif not articles_to_process:
            st.error("No articles to process")
        else:
            with st.spinner("Cleaning articles..."):
                # Process articles
                results = asyncio.run(process_articles(client, articles_to_process, output_language, article_summary_sentences, is_newsdesk))
                
                # Parse results
                cleaned_articles = []
                for i, result in enumerate(results):
                    try:
                        if is_newsdesk:
                            # For newsdesk, result should be a JSON array
                            parsed_data = json.loads(result)
                            
                            # Handle both formats: direct array or wrapped in "articles" key
                            if isinstance(parsed_data, list):
                                articles_array = parsed_data
                            elif isinstance(parsed_data, dict) and "articles" in parsed_data:
                                articles_array = parsed_data["articles"]
                            elif isinstance(parsed_data, dict):
                                articles_array = [parsed_data]
                            else:
                                continue
                            
                            cleaned_articles.extend(articles_array)
                        else:
                            # For regular articles, result is a single JSON object
                            article_data = json.loads(result)
                            cleaned_articles.append(article_data)
                            
                    except json.JSONDecodeError as e:
                        st.error(f"Error parsing response for {article_names[i]}: {str(e)}")
                    except Exception as e:
                        st.error(f"Unexpected error processing result {i+1}: {str(e)}")
                
                if cleaned_articles:
                    # Create temporary directory for this session if not exists
                    if not workdir:
                        workdir = tempfile.mkdtemp()
                    
                    # Create Word documents
                    doc_paths = []
                    for i, article in enumerate(cleaned_articles):
                        # Generate proper filename
                        source = article.get("source", "Unknown")
                        title = article.get("title", f"doc_{i}")
                        date = article.get("date", "No date")
                        
                        # Use the new filename generation
                        proper_filename = generate_docx_filename(source, title, date)
                        doc_path = os.path.join(workdir, proper_filename)
                        
                        # Assign images to articles (simple distribution for now)
                        article_images = all_images_info if include_images else None
                        
                        # The create_word_doc_with_images function will handle the proper naming
                        actual_path = create_word_doc_with_images(article, doc_path, article_images, is_newsdesk)
                        doc_paths.append(actual_path)
                    
                    # Save results to session state
                    st.session_state.cleaned_articles_tab1 = cleaned_articles
                    st.session_state.doc_paths_tab1 = doc_paths
                    st.session_state.workdir_tab1 = workdir
                    st.session_state.all_images_tab1 = all_images_info
                    
                    st.success(f"Successfully cleaned {len(cleaned_articles)} articles!")

    # Display results if available
    if 'cleaned_articles_tab1' in st.session_state:
        st.subheader("Cleaned Articles")
        
        # Show image info if available
        if 'all_images_tab1' in st.session_state and st.session_state.all_images_tab1:
            st.info(f"🖼️ {len(st.session_state.all_images_tab1)} images were extracted and included in the documents")
        
        # Display articles
        for i, article in enumerate(st.session_state.cleaned_articles_tab1):
            with st.expander(f"{article.get('title', f'Article {i+1}')} - {article.get('source', 'Unknown')}"):
                col1, col2 = st.columns([1, 1])
                with col1:
                    st.markdown(f"**Source:** {article.get('source', 'Unknown')}")
                    st.markdown(f"**Date:** {article.get('date', 'No date')}")
                with col2:
                    st.markdown(f"**Output Language:** {output_language}")
                    st.markdown(f"**Summary Length:** {article_summary_sentences} sentences")
                
                st.markdown("**Summary:**")
                st.write(article.get('summary', 'No summary'))
                
                st.markdown("**Cleaned & Translated Text:**")
                st.text_area("", article.get("cleaned_text", "No content"), height=200, key=f"clean_text_{i}")
        
        # Download options
        st.subheader("Download Options")
        
        col1, col2 = st.columns([1, 1])
        
        with col1:
            st.markdown("**Individual Downloads:**")
            for i, path in enumerate(st.session_state.doc_paths_tab1):
                article = st.session_state.cleaned_articles_tab1[i]
                st.markdown(
                    get_download_link(path, f"📄 {article.get('title', f'Article {i+1}')}"), 
                    unsafe_allow_html=True
                )
        
        with col2:
            if len(st.session_state.doc_paths_tab1) > 1:
                st.markdown("**Batch Download:**")
                zip_data = create_zip_from_files(st.session_state.doc_paths_tab1)
                b64_zip = base64.b64encode(zip_data).decode()
                href = f'<a href="data:application/zip;base64,{b64_zip}" download="cleaned_articles.zip">📦 Download All Articles (ZIP)</a>'
                st.markdown(href, unsafe_allow_html=True)

# Newsletter Generator Tab
with tab2:
    st.header("📰 Newsletter Generator")
    
    # Get language setting
    output_language = st.session_state.get('output_language', 'French')
    article_summary_sentences = st.session_state.get('article_summary_sentences', 3)
    exec_summary_sentences = st.session_state.get('exec_summary_sentences', 7)
    
    st.info(f"Newsletter will be generated in: {output_language} | Article summaries: ~{article_summary_sentences} sentences | Executive summary: ~{exec_summary_sentences} sentences")
    
    st.subheader("Article Sources")
    
    # Option to use cleaned articles from Tab 1 or upload new ones
    newsletter_source = st.radio(
        "Choose article source for newsletter:",
        ["Use cleaned articles from Article Cleaning tab", "Upload/Import new articles"],
        help="You can either use articles you've already cleaned, or process new articles specifically for the newsletter"
    )
    
    articles_for_newsletter = []
    newsletter_articles_to_process = []
    newsletter_article_names = []
    newsletter_images = []
    newsletter_workdir = None
        
    if newsletter_source == "Use cleaned articles from Article Cleaning tab":
        if 'cleaned_articles_tab1' in st.session_state:
            articles_for_newsletter = st.session_state.cleaned_articles_tab1
            newsletter_images = st.session_state.get('all_images_tab1', [])
            st.success(f"Found {len(articles_for_newsletter)} cleaned articles ready for newsletter generation")
            
            if newsletter_images:
                st.info(f"🖼️ {len(newsletter_images)} images will be included in the newsletter")
            
            # Show preview of articles
            with st.expander("Preview articles for newsletter"):
                for i, article in enumerate(articles_for_newsletter):
                    st.markdown(f"**{i+1}. {article.get('title', 'Untitled')}** - {article.get('source', 'Unknown')}")
        else:
            st.warning("No cleaned articles found. Please clean some articles in the 'Article Cleaning' tab first, or choose to upload new articles.")
    
    else:  # Upload/Import new articles
        st.subheader("Upload Articles for Newsletter")
        
        # Processing options for newsletter
        col1, col2 = st.columns([1, 1])
        with col1:
            newsletter_is_newsdesk = st.toggle(
                "📰 Newsdesk Mode",
                value=False,
                help="Enable this for Nexis Newsdesk documents",
                key="newsletter_newsdesk"
            )
        
        with col2:
            newsletter_include_images = st.toggle(
                "🖼️ Include Images",
                value=True,
                help="Extract and include images in the newsletter",
                key="newsletter_images"
            )
        
        # Input method for newsletter
        newsletter_input_method = st.radio(
            "Choose input method:",
            ["Upload Files", "Copy-Paste Text"],
            horizontal=True,
            key="newsletter_input"
        )
        
        if newsletter_input_method == "Upload Files":
            newsletter_uploaded_files = st.file_uploader(
                "Upload files for newsletter", 
                type=["pdf", "txt", "docx", "md"], 
                accept_multiple_files=True,
                key="newsletter_upload"
            )
            
            if newsletter_uploaded_files:
                paths, newsletter_workdir = get_paths_with_workdir(newsletter_uploaded_files)
                for path in paths:
                    if newsletter_include_images:
                        content, images = read_file_content_with_images(path, newsletter_workdir)
                        newsletter_images.extend(images)
                    else:
                        content = read_file_content(path)
                        
                    if content:
                        newsletter_articles_to_process.append(content)
                        newsletter_article_names.append(os.path.basename(path))
        
        else:  # Copy-Paste for newsletter
            newsletter_pasted_text = st.text_area(
                "Paste articles for newsletter (separate multiple articles with '---'):",
                height=300,
                placeholder="Copy and paste your article content here...\n\n---\n\nUse --- to separate multiple articles",
                key="newsletter_paste"
            )
            
            if newsletter_pasted_text.strip():
                # Split by --- if multiple articles
                split_articles = [art.strip() for art in newsletter_pasted_text.split('---') if art.strip()]
                newsletter_articles_to_process.extend(split_articles)
                newsletter_article_names.extend([f"Pasted_Article_{i+1}" for i in range(len(split_articles))])
        
        # Show what will be processed
        if newsletter_articles_to_process:
            col1, col2 = st.columns([2, 1])
            with col1:
                st.info(f"Ready to process {len(newsletter_articles_to_process)} articles for newsletter")
            with col2:
                if newsletter_images:
                    st.success(f"🖼️ Found {len(newsletter_images)} images")
                    
            with st.expander("Preview articles to process"):
                for i, name in enumerate(newsletter_article_names):
                    st.markdown(f"**{i+1}. {name}**")
    
    # Single button to generate newsletter
    generate_newsletter_button = False
    
    if newsletter_source == "Use cleaned articles from Article Cleaning tab":
        if articles_for_newsletter:
            generate_newsletter_button = st.button("📰 Generate Newsletter", use_container_width=True, key="gen_newsletter_existing")
    else:
        if newsletter_articles_to_process:
            generate_newsletter_button = st.button("🔄 Process Articles & Generate Newsletter", use_container_width=True, key="gen_newsletter_new")
    
    # Generate Newsletter (unified process)
    if generate_newsletter_button:
        client = get_client(key)
        if not client:
            st.error("Failed to initialize Azure OpenAI client")
        else:
            # If we need to process new articles first
            if newsletter_source != "Use cleaned articles from Article Cleaning tab":
                with st.spinner("Processing articles for newsletter..."):
                    # Process articles
                    results = asyncio.run(process_articles(client, newsletter_articles_to_process, output_language, article_summary_sentences, newsletter_is_newsdesk))
                    
                    # Parse results
                    cleaned_articles = []
                    for i, result in enumerate(results):
                        try:
                            if newsletter_is_newsdesk:
                                # For newsdesk, result should be a JSON array
                                parsed_data = json.loads(result)
                                if isinstance(parsed_data, list):
                                    articles_array = parsed_data
                                elif isinstance(parsed_data, dict) and "articles" in parsed_data:
                                    articles_array = parsed_data["articles"]
                                elif isinstance(parsed_data, dict):
                                    articles_array = [parsed_data]
                                else:
                                    continue
                                cleaned_articles.extend(articles_array)
                            else:
                                # For regular articles, result is a single JSON object
                                article_data = json.loads(result)
                                cleaned_articles.append(article_data)
                        except json.JSONDecodeError:
                            st.error(f"Error parsing response for {newsletter_article_names[i]}")
                    
                    if cleaned_articles:
                        articles_for_newsletter = cleaned_articles
                        st.success(f"Successfully processed {len(cleaned_articles)} articles!")
                    else:
                        st.error("Failed to process articles")
                        articles_for_newsletter = []
            
            # Generate newsletter if we have articles
            if articles_for_newsletter:
                with st.spinner("Generating newsletter..."):
                    # Extract summaries for newsletter
                    article_summaries = []
                    for article in articles_for_newsletter:
                        article_summaries.append({
                            'title': article.get('title', 'Untitled'),
                            'source': article.get('source', 'Unknown'),
                            'date': article.get('date', 'No date'),
                            'summary': article.get('summary', 'No summary available')
                        })
                    
                    # Generate executive summary
                    exec_summary = asyncio.run(generate_executive_summary(client, articles_for_newsletter, output_language, exec_summary_sentences))
                    
                    # Save to session state
                    st.session_state.article_summaries_newsletter = article_summaries
                    st.session_state.exec_summary_newsletter = exec_summary
                    st.session_state.newsletter_language = output_language
                    st.session_state.newsletter_article_sentences = article_summary_sentences
                    st.session_state.newsletter_exec_sentences = exec_summary_sentences
                    st.session_state.newsletter_images = newsletter_images
                    st.session_state.newsletter_workdir = newsletter_workdir or st.session_state.get('workdir_tab1')
                    
                    st.success("Newsletter generated successfully!")
    
    # Display Newsletter Preview
    if 'exec_summary_newsletter' in st.session_state:
        st.subheader("Newsletter Preview")
        
        # Language info
        newsletter_lang = st.session_state.get('newsletter_language', 'French')
        newsletter_images_count = len(st.session_state.get('newsletter_images', []))
        
        col1, col2 = st.columns([2, 1])
        with col1:
            st.info(f"Newsletter language: {newsletter_lang}")
        with col2:
            if newsletter_images_count > 0:
                st.success(f"🖼️ {newsletter_images_count} images included")
        
        # Executive Summary
        st.markdown("### Executive Summary")
        st.write(st.session_state.exec_summary_newsletter)
        st.markdown("---")
        
        # Article summaries
        st.markdown("### Articles")
        for article in st.session_state.article_summaries_newsletter:
            st.markdown(f"**{article['title']}** — *{article['source']}, {article['date']}*")
            st.write(article['summary'])
            st.markdown("---")
        
        # Download Newsletter
        if st.button("📄 Download Newsletter as Word Document", use_container_width=True):
            with st.spinner("Creating newsletter document..."):
                # Create temporary directory if it doesn't exist
                workdir = st.session_state.get('newsletter_workdir') or tempfile.mkdtemp()
                
                # Create newsletter document with images
                newsletter_filename = "newsletter.docx"
                newsletter_images = st.session_state.get('newsletter_images', [])
                
                doc_path = create_newsletter_doc_with_images(
                    st.session_state.exec_summary_newsletter, 
                    st.session_state.article_summaries_newsletter,
                    os.path.join(workdir, newsletter_filename),
                    newsletter_images
                )
                
                # Generate download link
                st.markdown(
                    get_download_link(doc_path, "📄 Download Newsletter Document"),
                    unsafe_allow_html=True
                )
                
                st.success("Newsletter document ready for download!")
    
    elif not articles_for_newsletter and newsletter_source == "Use cleaned articles from Article Cleaning tab":
        st.info("Please clean some articles in the 'Article Cleaning' tab first, or choose to upload new articles.")
    elif not newsletter_articles_to_process and newsletter_source != "Use cleaned articles from Article Cleaning tab":
        st.info("Please upload files or paste text to generate a newsletter.")

# Backward compatibility functions
def create_word_doc(article_data, output_path):
    """Create a Word document from cleaned article data using source-based template"""
    return create_word_doc_with_images(article_data, output_path, None, False)

def create_newsletter_doc(exec_summary, articles, output_path):
    """Create a Word document with newsletter content"""
    return create_newsletter_doc_with_images(exec_summary, articles, output_path, None)
